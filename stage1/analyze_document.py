#!/usr/bin/env python3
"""Analyze PDF and DOCX files to detect chapters and structure."""

import pymupdf as fitz
from docx import Document
import re
import sys
import os

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))        # .../stage1/
_COMBINED_DIR = os.path.dirname(_SCRIPT_DIR)                     # .../final-combined-pdf-to-lyx/
_PROJECT_ROOT = os.path.join(_COMBINED_DIR, 'input')             # .../input/  (PDF+DOCX drop zone)

def analyze_pdf(pdf_path):
    """Analyze PDF for chapters and text structure."""
    print(f"Analyzing PDF: {pdf_path}")
    print("=" * 70)

    doc = fitz.open(pdf_path)
    chapters = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        blocks = page.get_text("dict")["blocks"]

        for block in blocks:
            if "lines" in block:
                for line in block["lines"]:
                    for span in line["spans"]:
                        text = span["text"].strip()
                        size = round(span["size"], 1)
                        color = span["color"]
                        flags = span["flags"]  # bit 4 = bold
                        is_bold = bool(flags & 16)

                        # Detect chapter headings:
                        # - "Chapter Notes: <name>" at 19-21pt (EduRev format)
                        # - "Chapter <N>" numbered headings at any size
                        is_chapter_notes = (19 <= size <= 21 and
                                            re.search(r'chapter\s+notes\s*[:\-]', text, re.IGNORECASE))
                        is_numbered_chapter = re.search(r'chapter\s+\d+', text, re.IGNORECASE)
                        if text and (is_chapter_notes or is_numbered_chapter):
                            chapters.append({
                                'text': text,
                                'page': page_num + 1,
                                'size': size,
                                'color': color,
                                'bold': is_bold
                            })
                            print(f"Found chapter marker on page {page_num + 1}:")
                            print(f"  Text: '{text}'")
                            print(f"  Font size: {size}pt, Bold: {is_bold}, Color: {color}")

                        # Sample first few headings for detection rule verification
                        if page_num < 3 and size >= 13 and text:  # First 3 pages, large text
                            if len(text) < 100:  # Likely a heading, not body text
                                print(f"\nPage {page_num + 1} - Potential heading:")
                                print(f"  Text: '{text[:60]}...' " if len(text) > 60 else f"  Text: '{text}'")
                                print(f"  Font size: {size}pt, Bold: {is_bold}, Color: {color}")

    print(f"\n{'=' * 70}")
    print(f"Total pages: {len(doc)}")
    print(f"Chapters detected: {len(chapters)}")

    if chapters:
        print("\nChapter summary:")
        for ch in chapters:
            print(f"  - Page {ch['page']}: {ch['text']}")

    doc.close()
    return chapters

def analyze_docx(docx_path):
    """Analyze DOCX for structure and potential issues."""
    print(f"\n\nAnalyzing DOCX: {docx_path}")
    print("=" * 70)

    doc = Document(docx_path)

    # Count elements
    paragraphs = len(doc.paragraphs)
    tables = len(doc.tables)

    # Count images
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1

    # Check for headers/footers
    headers = sum(1 for section in doc.sections if section.header.paragraphs)
    footers = sum(1 for section in doc.sections if section.footer.paragraphs)

    # Sample some paragraphs
    print(f"Total paragraphs: {paragraphs}")
    print(f"Total tables: {tables}")
    print(f"Total images: {image_count}")
    print(f"Sections with headers: {headers}")
    print(f"Sections with footers: {footers}")

    # Check for edurev mentions
    edurev_count = 0
    sample_paragraphs = []

    for i, para in enumerate(doc.paragraphs[:20]):  # First 20 paragraphs
        text = para.text.strip()
        if text:
            if re.search(r'edurev', text, re.IGNORECASE):
                edurev_count += 1
                print(f"\n⚠️  Found 'edurev' in paragraph {i + 1}: '{text[:80]}...'")

            if len(sample_paragraphs) < 5 and len(text) > 20:
                sample_paragraphs.append((i + 1, text[:100]))

    print(f"\n'edurev' mentions found: {edurev_count}")

    if sample_paragraphs:
        print("\nSample paragraphs (first few with content):")
        for num, text in sample_paragraphs:
            print(f"  Para {num}: {text}...")

    return {
        'paragraphs': paragraphs,
        'tables': tables,
        'images': image_count,
        'edurev_mentions': edurev_count
    }

if __name__ == "__main__":
    if len(sys.argv) == 3:
        pdf_file = sys.argv[1]
        docx_file = sys.argv[2]
    elif len(sys.argv) == 2:
        # Single arg: base name (with or without extension)
        base = sys.argv[1]
        if base.endswith(".pdf"):
            base = base[:-4]
        pdf_file = os.path.join(_PROJECT_ROOT, f"{base}.pdf")
        docx_file = os.path.join(_PROJECT_ROOT, f"{base}.docx")
    else:
        # Auto-detect: find the only PDF/DOCX pair in project root
        pdfs = [f for f in os.listdir(_PROJECT_ROOT) if f.endswith(".pdf")]
        if len(pdfs) == 1:
            base = pdfs[0][:-4]
            pdf_file = os.path.join(_PROJECT_ROOT, pdfs[0])
            docx_file = os.path.join(_PROJECT_ROOT, f"{base}.docx")
        else:
            print("Usage: analyze_document.py [<base-name> | <file.pdf> <file.docx>]")
            print(f"PDFs found in project root: {pdfs}")
            sys.exit(1)

    if not os.path.exists(pdf_file):
        print(f"Error: PDF not found: {pdf_file}")
        sys.exit(1)
    if not os.path.exists(docx_file):
        print(f"Error: DOCX not found: {docx_file}")
        sys.exit(1)

    chapters = analyze_pdf(pdf_file)
    docx_info = analyze_docx(docx_file)

    print("\n" + "=" * 70)
    print("ANALYSIS SUMMARY")
    print("=" * 70)

    if len(chapters) == 0:
        print("✓ Single chapter document (no chapter markers detected)")
    elif len(chapters) == 1:
        print(f"✓ Single chapter: {chapters[0]['text']}")
    else:
        print(f"⚠️  Multiple chapters detected ({len(chapters)})")
        print("   User should confirm if they want single or split output")

    if docx_info['edurev_mentions'] > 0:
        print(f"⚠️  Document contains {docx_info['edurev_mentions']} 'edurev' references (will be removed)")

    print(f"\n📄 Document stats: {docx_info['paragraphs']} paragraphs, {docx_info['tables']} tables, {docx_info['images']} images")
