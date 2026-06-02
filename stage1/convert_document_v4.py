#!/usr/bin/env python3
"""Convert PDF to DOCX using extracted images from DOCX zip."""

import pymupdf as fitz
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from collections import deque
from PIL import Image
import re
import logging
import os
import glob
import zipfile
import tempfile
import shutil

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))        # .../stage1/
_COMBINED_DIR = os.path.dirname(_SCRIPT_DIR)                     # .../final-combined-pdf-to-lyx/
_PROJECT_ROOT = os.path.join(_COMBINED_DIR, 'input')             # .../input/  (PDF+DOCX drop zone)
_OUTPUT_DIR = os.path.join(_SCRIPT_DIR, 'output')               # .../stage1/output/
_LOGS_DIR = os.path.join(_SCRIPT_DIR, 'logs')                   # .../stage1/logs/
os.makedirs(_OUTPUT_DIR, exist_ok=True)
os.makedirs(_LOGS_DIR, exist_ok=True)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(os.path.join(_LOGS_DIR, 'conversion.log')),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

CONTENT_WIDTH_INCHES = 6.77  # A4 8.27" − 2×0.75" margins

# Bold-weight font name patterns — PDFs may not set flags & 16 for these
_BOLD_FONT_RE = re.compile(r'bold|black|heavy|semibold', re.IGNORECASE)

# Unicode ligature → ASCII replacements (common in PDF fonts)
_LIGATURES = {
    'ﬀ': 'ff',
    'ﬁ': 'fi',
    'ﬂ': 'fl',
    'ﬃ': 'ffi',
    'ﬄ': 'ffl',
    'ﬅ': 'st',
    'ﬆ': 'st',
}

class DocumentConverter:
    """Convert PDF to standardized DOCX format."""

    def __init__(self, pdf_path, docx_path, output_path):
        self.pdf_path = pdf_path
        self.docx_path = docx_path
        self.output_path = output_path
        self.image_files = []
        self.image_index = 0
        self.extract_dir = None

    def extract_images_from_docx_zip(self):
        """Extract images directly from DOCX zip file."""
        logger.info(f"Extracting images from DOCX zip: {self.docx_path}")

        self.extract_dir = tempfile.mkdtemp(prefix='pdf_conv_')
        with zipfile.ZipFile(self.docx_path, 'r') as zip_ref:
            zip_ref.extractall(self.extract_dir)

        # Get all image files from word/media, sorted by name
        media_path = os.path.join(self.extract_dir, 'word', 'media')
        if os.path.exists(media_path):
            images = glob.glob(os.path.join(media_path, 'image*'))
            # Sort by image number (image1, image2, etc.)
            images.sort(key=lambda x: int(''.join(filter(str.isdigit, os.path.basename(x)))))

            # Filter out very small images (likely watermarks) and EduRev logos
            filtered_images = []
            for img_path in images:
                size = os.path.getsize(img_path)
                basename = os.path.basename(img_path)
                _, ext = os.path.splitext(basename)

                if size < 1000:  # Less than 1KB - likely watermark
                    logger.info(f"Skipping tiny image (watermark): {basename} ({size} bytes)")
                    continue

                # EduRev logo: PNG file with banner-logo dimensions (~301×112)
                if ext.lower() == '.png':
                    try:
                        with Image.open(img_path) as im:
                            w, h = im.width, im.height
                        if 270 <= w <= 330 and 90 <= h <= 130:
                            logger.info(f"Skipping EduRev logo: {basename} ({w}×{h})")
                            continue
                    except Exception:
                        pass

                filtered_images.append(img_path)
                logger.debug(f"Using image: {basename} ({size} bytes)")

            self.image_files = filtered_images
            logger.info(f"Found {len(filtered_images)} usable images (filtered from {len(images)} total)")

        return self.image_files

    def sanitize_text(self, text):
        """Remove invalid XML characters from text."""
        if not text:
            return ""

        cleaned = ""
        for char in text:
            code = ord(char)
            if (32 <= code <= 126) or code in (9, 10, 13) or code >= 128:
                cleaned += char
            else:
                cleaned += " "

        cleaned = re.sub(r'\s+', ' ', cleaned).strip()
        return cleaned

    def _expand_ligatures_with_bold_map(self, raw_chars, bold_map):
        """Expand MuPDF ligature placeholders (control chars 0x01-0x05), preserving bold map.

        Returns (list[str], list[bool]) — expanded chars and their bold flags.
        Same three-rule heuristic as before; each ligature expands to 2 chars,
        both inheriting the placeholder's bold value.
        """
        chars, bolds = [], []
        for i, c in enumerate(raw_chars):
            if 0 < ord(c) < 32:
                prev = chars[-1].lower() if chars else ''
                nxt  = raw_chars[i + 1].lower() if i + 1 < len(raw_chars) else ''
                nxt2 = raw_chars[i + 2].lower() if i + 2 < len(raw_chars) else ''
                if prev == 'i':
                    lig = 'fi'
                elif nxt == 'e' and nxt2 == 'l':
                    lig = 'fi'
                elif nxt in 'aeiou':
                    lig = 'fl'
                else:
                    lig = 'fi'
                for ch in lig:
                    chars.append(ch)
                    bolds.append(bold_map[i])
            else:
                chars.append(c)
                bolds.append(bold_map[i])
        return chars, bolds

    def detect_heading_level(self, size, color, bold, text):
        """Detect heading level from PDF formatting."""
        # H1: 19–21pt (chapter titles at 19.5pt black, section titles at 21pt colored)
        if 19 <= size <= 21:
            return 1
        # H2: 16–17pt (covers 16.5pt used in newer documents)
        if 16 <= size <= 17:
            return 2
        if 13 <= size <= 14 and bold and len(text) < 80:
            return 3
        return 0

    def is_edurev_watermark(self, text, size):
        """Check if text is EDUREV watermark."""
        if re.search(r'edurev|durev', text, re.IGNORECASE):
            return True
        if size > 50:
            return True
        return False

    def is_toc_content(self, text):
        """Check if text is table of contents."""
        if re.search(r'table\s+of\s+contents', text, re.IGNORECASE):
            return True
        if re.search(r'\.{3,}\s*\d+$', text):
            return True
        return False

    def is_browser_artifact(self, text):
        """Check if text is a browser print artifact."""
        t = text.strip()
        if t == 'Firefox':
            return True
        # Page counters: "N of M" or "N/M"
        if re.match(r'^\d+\s+of\s+\d+$', t):
            return True
        if re.match(r'^\d+/\d+$', t):
            return True
        # Timestamps: DD/MM/YY or DD/MM/YYYY, HH:MM
        if re.match(r'^\d{2}/\d{2}/\d{2,4},\s+\d{2}:\d{2}$', t):
            return True
        # EduRev URLs
        if re.search(r'edurev\.in', text, re.IGNORECASE):
            return True
        # Browser tab title: "... Chapter Notes | Science Class N ..."
        if ' | ' in text and re.search(r'chapter notes|science class\s*\d+|pdf\s+down', text, re.IGNORECASE):
            return True
        return False

    def is_mcq_content(self, text, size, color, bold):
        """Check if text is part of an MCQ block."""
        # MCQ section label
        if re.search(r'multiple\s+choice\s+question', text, re.IGNORECASE):
            return True
        # "Try yourself:" — green bold label
        if re.search(r'try\s+yourself', text, re.IGNORECASE) and bold:
            return True
        # "View Solution" and "View More" — EduRev UI buttons
        if re.search(r'^view\s+(solution|more)$', text.strip(), re.IGNORECASE):
            return True
        # MCQ option labels: "A", "B", "C", "D" at ~9.8pt bold
        if re.match(r'^[A-D]$', text.strip()) and bold and size < 11:
            return True
        return False

    def extract_page_content(self, page, page_num, pdf_img_dims):
        """Extract text and image positions from PDF page."""
        content_items = []

        # Extract text blocks — use rawdict for char-level data (handles ligature placeholders)
        blocks = page.get_text("rawdict")["blocks"]

        in_mcq_block = False   # True between "Try yourself:" and "View Solution"
        after_option_label = False  # True right after A/B/C/D label (skip option text)
        in_toc_section = False  # True between "Table of Contents" and first H2/H1 content

        for block in blocks:
            if "lines" not in block:
                continue

            block_y = block["bbox"][1]

            for line in block["lines"]:
                raw_chars, bold_chars = [], []
                line_size = 0
                line_color = 0
                line_bold = False

                for span in line["spans"]:
                    is_bold = bool(span["flags"] & 16) or bool(_BOLD_FONT_RE.search(span.get("font", "")))
                    for ch in span["chars"]:
                        c = ch["c"]
                        if c.strip():
                            line_size  = round(span["size"], 1)
                            line_color = span["color"]
                            line_bold  = is_bold
                        raw_chars.append(c)
                        bold_chars.append(is_bold)

                # Expand ligature placeholders (preserving bold map)
                exp_chars, exp_bolds = self._expand_ligatures_with_bold_map(raw_chars, bold_chars)

                # Expand Unicode ligature chars (ﬁ → fi, ﬂ → fl, …)
                final_chars, final_bolds = [], []
                for c, b in zip(exp_chars, exp_bolds):
                    if c in _LIGATURES:
                        for ch in _LIGATURES[c]:
                            final_chars.append(ch)
                            final_bolds.append(b)
                    else:
                        final_chars.append(c)
                        final_bolds.append(b)

                line_text = re.sub(r'\s+', ' ', ''.join(final_chars)).strip()

                if not line_text:
                    continue
                if self.is_edurev_watermark(line_text, line_size):
                    continue

                # TOC state machine: track when we're inside a Table of Contents
                if self.is_toc_content(line_text):
                    in_toc_section = True
                    continue
                # Numbered TOC entries: "1. Heading text" — skip while in TOC section
                if in_toc_section and re.match(r'^\d+[\.\)]\s+\S', line_text):
                    continue
                # Exit TOC mode when a real content heading (H2+) appears
                if in_toc_section and line_size >= 16:
                    in_toc_section = False

                if self.is_browser_artifact(line_text):
                    logger.debug(f"Skipping browser artifact: {line_text}")
                    continue

                # MCQ state machine
                if re.search(r'try\s+yourself', line_text, re.IGNORECASE) and line_bold:
                    in_mcq_block = True
                    logger.debug(f"MCQ block start: {line_text}")
                    continue

                if re.search(r'^view\s+solution$', line_text.strip(), re.IGNORECASE):
                    in_mcq_block = False
                    after_option_label = False
                    logger.debug("MCQ block end")
                    continue

                if self.is_mcq_content(line_text, line_size, line_color, line_bold):
                    if re.match(r'^[A-D]$', line_text.strip()) and line_bold and line_size < 11:
                        after_option_label = True
                    logger.debug(f"Skipping MCQ content: {line_text}")
                    continue

                if in_mcq_block:
                    logger.debug(f"Skipping MCQ question text: {line_text}")
                    continue

                if after_option_label:
                    logger.debug(f"Skipping MCQ option text: {line_text}")
                    after_option_label = False
                    continue

                # Group final_chars into (text_segment, is_bold) runs.
                # Use re.sub without strip() so boundary spaces are preserved
                # (sanitize_text strips, which drops the space between runs).
                runs = []
                if final_chars:
                    seg, seg_bold = final_chars[0], final_bolds[0]
                    for c, b in zip(final_chars[1:], final_bolds[1:]):
                        if b == seg_bold:
                            seg += c
                        else:
                            cleaned = re.sub(r'\s+', ' ', seg)
                            if cleaned.strip():
                                runs.append((cleaned, seg_bold))
                            seg, seg_bold = c, b
                    cleaned = re.sub(r'\s+', ' ', seg)
                    if cleaned.strip():
                        runs.append((cleaned, seg_bold))

                if not runs:
                    continue

                content_items.append({
                    'type':       'text',
                    'y_position': block_y,
                    'runs':       runs,
                    'size':       line_size,
                    'color':      line_color,
                    'bold':       line_bold,
                })

        # Extract image positions
        images = page.get_images()
        for img in images:
            try:
                xref = img[0]
                img_rects = page.get_image_rects(xref)
                if img_rects:
                    rect = img_rects[0]
                    if rect.width > 50 and rect.height > 50:
                        px_w, px_h, pdf_pt_w = pdf_img_dims.get((page_num - 1, xref), (0, 0, 0))
                        content_items.append({
                            'type': 'image',
                            'y_position': rect.y0,
                            'xref': xref,
                            'px_w': px_w,
                            'px_h': px_h,
                            'pdf_pt_w': pdf_pt_w,
                        })
            except:
                pass

        # Sort by position
        content_items.sort(key=lambda x: x['y_position'])
        return content_items

    def _scan_pdf_image_slots(self, pdf_doc):
        """Pre-scan all PDF images: detect splits and build a dims lookup dict.

        Returns (split_set, dims_dict, active_slots_ordered) where:
          active_slots_ordered: list of (pnum, xref) in document reading order,
          excluding split/logo slots. Used to build a 1:1 mapping with DOCX images.
        """
        split_set = set()   # (page_num, xref) pairs that are split continuations
        dims_dict = {}      # (page_num, xref) → (px_w, px_h, pdf_pt_w)
        prev_dims = None
        prev_page = -1
        # Collect (pnum, y0, xref) for active slots; sort by (pnum, y0) for document order
        _raw_slots = []

        for pnum in range(len(pdf_doc)):
            page = pdf_doc[pnum]
            for img in page.get_images():
                xref = img[0]
                try:
                    rects = page.get_image_rects(xref)
                    if not rects:
                        continue
                    r = rects[0]
                    if r.width <= 50 or r.height <= 50:
                        continue
                    base = pdf_doc.extract_image(xref)
                    px_w = base.get('width', 0)
                    px_h = base.get('height', 0)
                    curr_dims = (px_w, px_h)
                    dims_dict[(pnum, xref)] = (px_w, px_h, r.width)

                    # EduRev logo: banner-shaped image near page top
                    if 270 <= px_w <= 330 and 90 <= px_h <= 130:
                        split_set.add((pnum, xref))
                        logger.info(f"Skipping EduRev logo slot: page {pnum+1}, xref {xref} ({px_w}×{px_h})")
                        continue

                    # Split: identical pixel dims as previous slot AND on the very next page
                    if curr_dims == prev_dims and pnum == prev_page + 1:
                        split_set.add((pnum, xref))
                        logger.info(f"Split continuation detected: page {pnum+1}, xref {xref} ({px_w}×{px_h})")
                    else:
                        prev_dims = curr_dims
                        prev_page = pnum
                        _raw_slots.append((pnum, r.y0, xref))
                except:
                    pass

        # Sort by page then y-position to match content_items reading order
        _raw_slots.sort(key=lambda x: (x[0], x[1]))
        active_slots_ordered = [(pnum, xref) for pnum, _y0, xref in _raw_slots]

        return split_set, dims_dict, active_slots_ordered

    def _get_docx_img_dims(self, img_path):
        """Get pixel dimensions of a DOCX image file."""
        try:
            with Image.open(img_path) as img:
                return (img.width, img.height)
        except:
            return (0, 0)

    def _is_dim_match(self, docx_dims, pdf_dims, tolerance=0.20):
        """True if DOCX dims match PDF dims exactly or share the same aspect ratio.

        Tolerance of 20% handles JPEG vs PNG re-encoding differences in this
        document type, where the same image may differ by up to ~16% in aspect.
        """
        dw, dh = docx_dims
        pw, ph = pdf_dims
        if dw == pw and dh == ph:
            return True
        if dh == 0 or ph == 0:
            return False
        return abs(dw / dh - pw / ph) < tolerance

    def insert_image(self, doc, image_path, width=None):
        """Insert image with NO spacing."""
        try:
            para = doc.add_paragraph()
            run = para.add_run()
            run.add_picture(image_path, width=width if width is not None else Inches(5.5))
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # CRITICAL: Set spacing to ZERO (no white lines)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0

            logger.debug(f"Inserted image {self.image_index + 1}: {os.path.basename(image_path)}")
            self.image_index += 1

        except Exception as e:
            logger.error(f"Failed to insert image: {e}")

    def add_text_paragraph(self, doc, runs, size, color, bold):
        """Add text paragraph with formatting. runs = [(text, is_bold), ...]"""
        heading_level = self.detect_heading_level(
            size, color, bold, ''.join(t for t, _ in runs))

        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        pt_size = Pt(16 if heading_level == 1 else 14 if heading_level == 2 else 12)
        for run_text, run_bold in runs:
            run = para.add_run(run_text)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.size = pt_size
            # Headings always fully bold; body text uses per-run bold from PDF
            run.font.bold = True if heading_level > 0 else run_bold

    def convert(self):
        """Main conversion process."""
        logger.info("=" * 70)
        logger.info("Starting conversion with extracted DOCX images")
        logger.info("=" * 70)

        try:
            # Extract images from DOCX zip
            self.extract_images_from_docx_zip()

            # Create new document
            doc = Document()

            # Set page layout
            for section in doc.sections:
                section.page_height = Inches(11.69)
                section.page_width = Inches(8.27)
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)

            # Process PDF
            pdf_doc = fitz.open(self.pdf_path)
            logger.info(f"Processing {len(pdf_doc)} pages")

            # Pre-scan: detect split continuations and build pixel-dims lookup
            split_set, pdf_img_dims, active_slots_ordered = self._scan_pdf_image_slots(pdf_doc)
            logger.info(f"Split continuations to skip: {len(split_set)}")
            logger.info(f"Active PDF image slots: {len(active_slots_ordered)}")

            # Build 1:1 assignment: each active PDF slot gets the Nth DOCX image.
            # This guarantees all DOCX images are spread evenly throughout the document
            # regardless of dimension differences between PDF and DOCX image versions.
            docx_imgs_list = self.image_files  # already filtered and sorted
            img_assignment = {}  # (pnum, xref) -> docx_image_path
            for i, (pnum, xref) in enumerate(active_slots_ordered):
                if i < len(docx_imgs_list):
                    img_assignment[(pnum, xref)] = docx_imgs_list[i]
                    logger.debug(f"Assigned slot page {pnum+1} xref={xref} → {os.path.basename(docx_imgs_list[i])}")
            trailing_imgs = docx_imgs_list[len(active_slots_ordered):]
            logger.info(f"DOCX images assigned: {len(img_assignment)}, trailing: {len(trailing_imgs)}")

            for page_num in range(len(pdf_doc)):
                logger.info(f"Page {page_num + 1}/{len(pdf_doc)}")
                page = pdf_doc[page_num]

                content_items = self.extract_page_content(page, page_num + 1, pdf_img_dims)

                for item in content_items:
                    if item['type'] == 'text':
                        self.add_text_paragraph(
                            doc,
                            item['runs'],
                            item['size'],
                            item['color'],
                            item['bold']
                        )
                    elif item['type'] == 'image':
                        if (page_num, item['xref']) in split_set:
                            logger.debug(f"Skipping split: page {page_num+1}, xref {item['xref']}")
                            continue
                        img_path = img_assignment.get((page_num, item['xref']))
                        if img_path:
                            pdf_in = item['pdf_pt_w'] / 72.0 if item.get('pdf_pt_w') else 0
                            img_width = Inches(min(pdf_in, CONTENT_WIDTH_INCHES)) if pdf_in > 0 else Inches(5.5)
                            self.insert_image(doc, img_path, width=img_width)

            pdf_doc.close()

            # Append any DOCX images beyond the active PDF slot count
            for img_path in trailing_imgs:
                logger.info(f"Appending trailing image: {os.path.basename(img_path)}")
                self.insert_image(doc, img_path, width=Inches(CONTENT_WIDTH_INCHES))

            # Save
            doc.save(self.output_path)

            logger.info("=" * 70)
            logger.info(f"✓ Complete! Images used: {self.image_index}/{len(self.image_files)}")
            logger.info("=" * 70)

        finally:
            if self.extract_dir and os.path.exists(self.extract_dir):
                shutil.rmtree(self.extract_dir)
                logger.debug("Cleaned up temp extraction dir")

        return self.output_path

    def _make_doc(self):
        """Create a new Document with standard A4 page layout."""
        doc = Document()
        for section in doc.sections:
            section.page_height = Inches(11.69)
            section.page_width = Inches(8.27)
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        return doc

    def _process_image_item(self, item, page_num, doc, split_set, docx_queue):
        """Insert the DOCX image corresponding to a PDF image slot (or skip if split/logo)."""
        if (page_num, item['xref']) in split_set:
            logger.debug(f"Skipping split/logo: page {page_num+1}, xref {item['xref']}")
            return

        pdf_dims = (item['px_w'], item['px_h'])

        _MAX_LOOKAHEAD = 8
        _matched_idx = None
        for _i, (_qp, _qd) in enumerate(list(docx_queue)[:_MAX_LOOKAHEAD]):
            if self._is_dim_match(_qd, pdf_dims):
                _matched_idx = _i
                break

        if _matched_idx is not None:
            for _ in range(_matched_idx):
                orphan_path, _ = docx_queue.popleft()
                logger.info(f"Inserting orphan: {os.path.basename(orphan_path)}")
                self.insert_image(doc, orphan_path, width=Inches(CONTENT_WIDTH_INCHES))
            img_path, _ = docx_queue.popleft()
            pdf_in = item['pdf_pt_w'] / 72.0 if item.get('pdf_pt_w') else 0
            img_width = Inches(min(pdf_in, CONTENT_WIDTH_INCHES)) if pdf_in > 0 else Inches(5.5)
            self.insert_image(doc, img_path, width=img_width)
        elif docx_queue:
            img_path, _ = docx_queue.popleft()
            logger.info(f"Sequential fallback for PDF trigger page {page_num+1} ({pdf_dims[0]}×{pdf_dims[1]}): inserting {os.path.basename(img_path)}")
            pdf_in = item['pdf_pt_w'] / 72.0 if item.get('pdf_pt_w') else 0
            img_width = Inches(min(pdf_in, CONTENT_WIDTH_INCHES)) if pdf_in > 0 else Inches(5.5)
            self.insert_image(doc, img_path, width=img_width)

    def convert_chapters(self, chapters):
        """Convert a multi-chapter PDF into separate DOCX files.

        chapters: list of dicts with keys:
          name       - chapter name (used in log messages)
          start_page - first page (1-based, inclusive)
          end_page   - last page (1-based, inclusive)
          output_path - where to save this chapter's DOCX
        """
        logger.info("=" * 70)
        logger.info(f"Starting chapter split conversion ({len(chapters)} chapters)")
        logger.info("=" * 70)

        try:
            self.extract_images_from_docx_zip()

            pdf_doc = fitz.open(self.pdf_path)
            logger.info(f"PDF: {len(pdf_doc)} pages")

            split_set, pdf_img_dims, _active_slots = self._scan_pdf_image_slots(pdf_doc)
            logger.info(f"Split/logo slots to skip: {len(split_set)}")

            docx_queue = deque(
                (img_path, self._get_docx_img_dims(img_path))
                for img_path in self.image_files
            )

            # Build page→chapter lookup
            page_to_ch = {}
            for ch in chapters:
                for p in range(ch['start_page'], ch['end_page'] + 1):
                    page_to_ch[p] = ch['name']

            # Create one Document per chapter
            ch_docs = {ch['name']: self._make_doc() for ch in chapters}

            for page_num in range(len(pdf_doc)):
                page_1 = page_num + 1
                ch_name = page_to_ch.get(page_1)
                if ch_name is None:
                    continue

                doc = ch_docs[ch_name]
                logger.info(f"Page {page_1}/{len(pdf_doc)} → {ch_name}")
                page = pdf_doc[page_num]
                content_items = self.extract_page_content(page, page_1, pdf_img_dims)

                for item in content_items:
                    if item['type'] == 'text':
                        self.add_text_paragraph(doc, item['runs'], item['size'], item['color'], item['bold'])
                    elif item['type'] == 'image':
                        self._process_image_item(item, page_num, doc, split_set, docx_queue)

            pdf_doc.close()

            # Flush any remaining DOCX images into the last chapter
            last_doc = ch_docs[chapters[-1]['name']]
            while docx_queue:
                orphan_path, _ = docx_queue.popleft()
                logger.info(f"Trailing orphan → last chapter: {os.path.basename(orphan_path)}")
                self.insert_image(last_doc, orphan_path, width=Inches(CONTENT_WIDTH_INCHES))

            # Save each chapter
            output_paths = []
            for ch in chapters:
                ch_docs[ch['name']].save(ch['output_path'])
                logger.info(f"✓ Saved: {ch['output_path']}")
                output_paths.append(ch['output_path'])

            logger.info("=" * 70)
            logger.info(f"✓ Done! {len(chapters)} files, {self.image_index} images total")
            logger.info("=" * 70)
            return output_paths

        finally:
            if self.extract_dir and os.path.exists(self.extract_dir):
                shutil.rmtree(self.extract_dir)
                logger.debug("Cleaned up temp extraction dir")


if __name__ == "__main__":
    import sys

    if len(sys.argv) >= 2:
        base = sys.argv[1]
        if base.endswith(".pdf"):
            base = base[:-4]
        pdf_path = os.path.join(_PROJECT_ROOT, f"{base}.pdf")
        docx_path = os.path.join(_PROJECT_ROOT, f"{base}.docx")
        output_path = os.path.join(_OUTPUT_DIR, f"{base}-formatted.docx")
    else:
        # Auto-detect: find the only PDF in project root
        pdfs = [f for f in os.listdir(_PROJECT_ROOT) if f.endswith(".pdf")]
        if len(pdfs) == 1:
            base = pdfs[0][:-4]
            pdf_path = os.path.join(_PROJECT_ROOT, pdfs[0])
            docx_path = os.path.join(_PROJECT_ROOT, f"{base}.docx")
            output_path = os.path.join(_OUTPUT_DIR, f"{base}-formatted.docx")
        else:
            print("Usage: convert_document_v4.py [<base-name>]")
            print(f"PDFs found in project root: {pdfs}")
            sys.exit(1)

    if not os.path.exists(pdf_path):
        print(f"Error: PDF not found: {pdf_path}")
        sys.exit(1)
    if not os.path.exists(docx_path):
        print(f"Error: DOCX not found: {docx_path}")
        sys.exit(1)

    converter = DocumentConverter(
        pdf_path=pdf_path,
        docx_path=docx_path,
        output_path=output_path
    )
    output_file = converter.convert()
    print(f"\n✓ Output: {output_file}")
